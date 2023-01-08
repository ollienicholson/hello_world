# Import the required libraries
import docx

# Open a new document
document = docx.Document()

# Add a heading
document.add_heading('Heading 1', 0)

document.add_heading('Heading 2', 1)
document.add_heading('Heading 3', 2)
document.add_heading('Heading 4', 3)
document.add_heading('Heading 5', 5)

# Add some text
document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed euismod nunc id dui tempus, a porttitor nisi vehicula. Fusce vel lacus vel velit tempor condimentum. Curabitur laoreet purus a tortor bibendum, non imperdiet dolor volutpat. Aenean eget nibh hendrerit, fringilla magna a, placerat quam. Sed mollis vehicula elit. Duis suscipit est id diam viverra, in tincidunt nibh aliquam. Curabitur malesuada euismod ornare. Suspendisse potenti.')

# Save the document
document.save('text_to_docx.docx')

print('\nDocument saved successfully!')
